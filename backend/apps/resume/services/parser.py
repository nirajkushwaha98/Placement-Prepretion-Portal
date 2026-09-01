import io
import os
import pdfplumber
import docx

class ResumeTextExtractor:
    @staticmethod
    def extract_text(file_obj, file_name: str) -> str:
        """
        Extracts clean UTF-8 text from PDF or DOCX file stream.
        """
        ext = os.path.splitext(file_name)[1].lower()

        if ext == '.pdf':
            return ResumeTextExtractor._extract_from_pdf(file_obj)
        elif ext in ('.docx', '.doc'):
            return ResumeTextExtractor._extract_from_docx(file_obj)
        elif ext == '.txt':
            try:
                return file_obj.read().decode('utf-8', errors='ignore')
            except Exception:
                return ""
        else:
            # Try reading as text
            try:
                return file_obj.read().decode('utf-8', errors='ignore')
            except Exception:
                return ""

    @staticmethod
    def _extract_from_pdf(file_obj) -> str:
        text_chunks = []
        try:
            # Check if file_obj is a file path or file-like object
            if hasattr(file_obj, 'seek'):
                file_obj.seek(0)

            with pdfplumber.open(file_obj) as pdf:
                for page in pdf.pages:
                    extracted = page.extract_text(layout=True) or page.extract_text()
                    if extracted:
                        text_chunks.append(extracted)
        except Exception as e:
            # Fallback to direct raw stream read if pdfplumber encounters error
            try:
                if hasattr(file_obj, 'seek'):
                    file_obj.seek(0)
                from pdfminer.high_level import extract_text as pdfminer_extract
                return pdfminer_extract(file_obj)
            except Exception:
                return ""

        return "\n".join(text_chunks).strip()

    @staticmethod
    def _extract_from_docx(file_obj) -> str:
        try:
            if hasattr(file_obj, 'seek'):
                file_obj.seek(0)
            doc = docx.Document(file_obj)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text.strip())
            return "\n".join(full_text).strip()
        except Exception:
            return ""
